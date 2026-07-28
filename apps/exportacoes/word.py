"""
Exportação da grade para Word (Módulo 14), usando python-docx.
"""
from io import BytesIO

from docx import Document


def gerar_docx_grade(titulo, dias, grid, montar_linhas):
    documento = Document()
    documento.add_heading(titulo, level=1)

    horarios = list(grid.keys())
    tabela = documento.add_table(rows=len(horarios) + 1, cols=len(dias) + 1)
    tabela.style = 'Table Grid'

    tabela.cell(0, 0).text = 'Horário'
    for indice, (_codigo, rotulo) in enumerate(dias, start=1):
        tabela.cell(0, indice).text = rotulo

    for linha_idx, horario in enumerate(horarios, start=1):
        tabela.cell(linha_idx, 0).text = f'{horario.inicio.strftime("%H:%M")}–{horario.fim.strftime("%H:%M")}'
        por_dia = grid[horario]
        for coluna_idx, (codigo, _rotulo) in enumerate(dias, start=1):
            aula = por_dia.get(codigo)
            tabela.cell(linha_idx, coluna_idx).text = '\n'.join(montar_linhas(aula)) if aula else ''

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer
